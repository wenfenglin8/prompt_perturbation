from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "results" / "errica_methodology_adaptation_for_proposal.docx"


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_formula(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    xml = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        "<m:r>"
        "<m:rPr><m:sty m:val=\"p\"/></m:rPr>"
        f"<m:t>{escape(text)}</m:t>"
        "</m:r>"
        "</m:oMath>"
    )
    paragraph._p.append(parse_xml(xml))


def add_numbered_formula(doc: Document, number: int, formula: str, meaning: str) -> None:
    add_formula(doc, f"({number})  {formula}")
    if meaning:
        doc.add_paragraph(f"Meaning: {meaning}")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()
    doc.add_heading("Adapting Errica et al.'s Methodology for the Proposal", level=0)

    doc.add_heading("1. Core Idea", level=1)
    doc.add_paragraph(
        "The methodology of What Did I Do Wrong? can be borrowed, but it should be adapted rather than copied directly. "
        "Errica et al. focus on sensitivity and consistency for classification prompts. The current proposal can extend this idea "
        "from label changes to generative-output semantic drift and correctness drift."
    )
    add_code_block(
        doc,
        "Errica et al.: prompt rephrasing -> prediction label changes\n"
        "This proposal: prompt perturbation -> output semantic distribution shift -> correctness drift",
    )

    doc.add_heading("2. Clean Formula Notation", level=1)
    doc.add_paragraph(
        "The formula notation is organized as follows. This section is the recommended version to copy into the proposal."
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Symbol"
    hdr[1].text = "Meaning"
    notation = [
        ("i", "index of a task instance or case"),
        ("k", "index of a prompt perturbation type"),
        ("r, q", "indices of repeated samples from the same prompt"),
        ("R", "number of repeated samples per prompt version"),
        ("p_i", "original prompt for case i"),
        ("T_k(p_i)", "prompt p_i after perturbation type k"),
        ("M(p_i; s_r)", "model output from prompt p_i under sampling randomness s_r"),
        ("z^0_{i,r}", "embedding of the r-th output from the original prompt"),
        ("z^k_{i,r}", "embedding of the r-th output from the perturbed prompt"),
        ("d(a,b)", "semantic distance between two embedded outputs"),
        ("D_raw(i,k)", "raw semantic drift between original and perturbed outputs"),
        ("N(i,k)", "within-prompt sample-noise baseline"),
        ("S_sem(i,k)", "noise-corrected semantic sensitivity"),
        ("C_0(i), C_k(i)", "mean correctness under original and perturbed prompts"),
        ("Delta C(i,k)", "absolute correctness drift"),
    ]
    for symbol, meaning in notation:
        cells = table.add_row().cells
        cells[0].text = symbol
        cells[1].text = meaning

    doc.add_heading("3. Recommended Core Formula Set", level=1)
    add_numbered_formula(
        doc,
        1,
        "d(a,b) = 1 - cosine_similarity(a,b)",
        "Defines semantic distance in the embedding space.",
    )
    add_numbered_formula(
        doc,
        2,
        "D_raw(i,k) = (1 / R²) ΣᵣΣᵩ d(z⁰ᵢ,ᵣ, zᵏᵢ,ᵩ)",
        "Measures the observed semantic distance between all original-prompt outputs and all perturbed-prompt outputs.",
    )
    add_numbered_formula(
        doc,
        3,
        "N₀(i) = [2 / R(R - 1)] Σᵣ<ᵩ d(z⁰ᵢ,ᵣ, z⁰ᵢ,ᵩ)",
        "Measures natural output variation when the original prompt is sampled repeatedly.",
    )
    add_numbered_formula(
        doc,
        4,
        "Nₖ(i) = [2 / R(R - 1)] Σᵣ<ᵩ d(zᵏᵢ,ᵣ, zᵏᵢ,ᵩ)",
        "Measures natural output variation when the perturbed prompt is sampled repeatedly.",
    )
    add_numbered_formula(
        doc,
        5,
        "N(i,k) = 1/2 [N₀(i) + Nₖ(i)]",
        "Combines original and perturbed within-prompt variation into a sample-noise baseline.",
    )
    add_numbered_formula(
        doc,
        6,
        "S_sem(i,k) = max(0, D_raw(i,k) - N(i,k))",
        "Defines the main proposed metric: noise-corrected semantic sensitivity.",
    )
    add_numbered_formula(
        doc,
        7,
        "C₀(i) = (1/R) Σᵣ c⁰ᵢ,ᵣ;    Cₖ(i) = (1/R) Σᵣ cᵏᵢ,ᵣ",
        "Defines repeated-sampling correctness for original and perturbed prompts.",
    )
    add_numbered_formula(
        doc,
        8,
        "ΔC(i,k) = |C₀(i) - Cₖ(i)|",
        "Measures the magnitude of correctness drift.",
    )
    add_numbered_formula(
        doc,
        9,
        "ρ = Spearman({S_sem(i,k)}, {ΔC(i,k)})",
        "Tests whether semantic sensitivity ranks align with correctness-drift ranks.",
    )
    add_numbered_formula(
        doc,
        10,
        "R(i,k) = S_sem(i,k) · ΔC(i,k)",
        "Optional composite robustness-risk score.",
    )
    doc.add_paragraph(
        "Equations (1)-(6) define noise-corrected semantic sensitivity. Equations (7)-(9) connect semantic sensitivity to correctness drift. Equation (10) is optional and can be presented as a proposed robustness-risk score."
    )

    doc.add_heading("4. From Classification Sensitivity to Generative Semantic Sensitivity", level=1)
    doc.add_paragraph("Let:")
    add_code_block(
        doc,
        "x_i = the i-th task instance\n"
        "p_i = original prompt\n"
        "T_k(p_i) = prompt after perturbation type k\n"
        "M(p_i; s_r) = model output under sampling randomness s_r",
    )
    doc.add_paragraph("Embed each output into a vector space:")
    add_code_block(
        doc,
        "z^0_{i,r} = embed(M(p_i; s_r))\n"
        "z^k_{i,r} = embed(M(T_k(p_i); s_r))",
    )
    doc.add_paragraph("Define semantic distance:")
    add_code_block(doc, "d(a, b) = 1 - cosine_similarity(a, b)")
    doc.add_paragraph("Raw cross-prompt drift:")
    add_formula(doc, "D_raw(i,k) = (1 / R²) Σᵣ Σᵩ d(z⁰ᵢ,ᵣ, zᵏᵢ,ᵩ)")
    doc.add_paragraph(
        "This generalizes classification sensitivity from whether a predicted label changes to how much the generated output distribution shifts in semantic space."
    )

    doc.add_heading("5. Core Innovation: Sample-Noise Correction", level=1)
    doc.add_paragraph(
        "Because generative LLMs are stochastic, output variation is not caused only by prompt perturbation. The same prompt can produce different outputs across repeated samples. "
        "Therefore, observed drift should be decomposed into perturbation-induced drift and sampling noise."
    )
    add_formula(doc, "observed sensitivity = perturbation-induced drift + sampling noise")
    doc.add_paragraph("Within-prompt sampling noise for the original prompt:")
    add_formula(doc, "N₀(i) = [2 / R(R - 1)] Σᵣ<ᵩ d(z⁰ᵢ,ᵣ, z⁰ᵢ,ᵩ)")
    doc.add_paragraph("Within-prompt sampling noise for the perturbed prompt:")
    add_formula(doc, "Nₖ(i) = [2 / R(R - 1)] Σᵣ<ᵩ d(zᵏᵢ,ᵣ, zᵏᵢ,ᵩ)")
    doc.add_paragraph("Combined sample-noise baseline:")
    add_formula(doc, "N(i,k) = 1/2 [N₀(i) + Nₖ(i)]")
    doc.add_paragraph("Noise-corrected semantic sensitivity:")
    add_formula(doc, "S_sem(i,k) = max(0, D_raw(i,k) - N(i,k))")
    doc.add_paragraph(
        "This is the most important methodological extension. It subtracts within-prompt sampling variability from cross-prompt semantic drift, making the metric better suited to stochastic generative LLMs."
    )

    doc.add_heading("6. From Consistency to Output Invariance", level=1)
    doc.add_paragraph(
        "Errica et al.'s consistency concept can be generalized from label-level consistency to semantic-output invariance."
    )
    add_formula(doc, "I_sem(i,k) = 1 - S_sem(i,k)")
    doc.add_paragraph("A smoother normalized version is:")
    add_formula(doc, "I_sem(i,k) = exp(-λ S_sem(i,k)), where λ > 0")
    add_bullets(
        doc,
        [
            "Large S_sem means high semantic sensitivity.",
            "Large I_sem means strong output invariance.",
            "Small I_sem means the model is unstable under the perturbation.",
        ],
    )

    doc.add_heading("7. Task-Level and Perturbation-Level Sensitivity", level=1)
    doc.add_paragraph("Average sensitivity for perturbation type k:")
    add_formula(doc, "Sₖ = (1 / |I|) Σᵢ S_sem(i,k)")
    doc.add_paragraph("Average sensitivity for task t:")
    add_formula(doc, "Sₜ = [1 / (|Iₜ||K|)] Σᵢ∈Iₜ Σₖ S_sem(i,k)")
    doc.add_paragraph("Task-by-perturbation sensitivity:")
    add_formula(doc, "Sₜ,ₖ = (1 / |Iₜ|) Σᵢ∈Iₜ S_sem(i,k)")
    doc.add_paragraph("This directly supports tornado charts and ANOVA:")
    add_formula(doc, "S_sem(i,k) = μ + αₜ + βₖ + (αβ)ₜ,ₖ + εᵢ,ₖ")

    doc.add_heading("8. Correctness Transfer: From Semantic Instability to Functional Instability", level=1)
    doc.add_paragraph("Define per-sample correctness:")
    add_code_block(doc, "c^0_{i,r} in {0,1}\nc^k_{i,r} in {0,1}")
    doc.add_paragraph("Repeated-sampling correctness:")
    add_formula(doc, "C₀(i) = (1/R) Σᵣ c⁰ᵢ,ᵣ")
    add_formula(doc, "Cₖ(i) = (1/R) Σᵣ cᵏᵢ,ᵣ")
    doc.add_paragraph("Correctness drift:")
    add_formula(doc, "ΔC(i,k) = |C₀(i) - Cₖ(i)|")
    doc.add_paragraph("Harmful correctness drop:")
    add_formula(doc, "H(i,k) = 1[C₀(i) > Cₖ(i)]")
    doc.add_paragraph("Then test whether semantic sensitivity predicts correctness drift:")
    add_formula(doc, "ρ = Spearman({S_sem(i,k)}, {ΔC(i,k)})")
    doc.add_paragraph("Or use a regression-style formulation:")
    add_formula(doc, "ΔC(i,k) = γ₀ + γ₁S_sem(i,k) + γ₂taskᵢ + γ₃perturbationₖ + εᵢ,ₖ")
    doc.add_paragraph("If gamma_1 > 0 or rho > 0, semantic sensitivity can be interpreted as an indicator of correctness drift.")

    doc.add_heading("9. Optional New Composite Metric: Noise-Corrected Robustness Risk", level=1)
    doc.add_paragraph("A composite robustness risk metric can be defined as:")
    add_formula(doc, "R(i,k) = S_sem(i,k) · ΔC(i,k)")
    doc.add_paragraph("A harmful-risk version is:")
    add_formula(doc, "R_harm(i,k) = S_sem(i,k) · H(i,k)")
    doc.add_paragraph("Task-by-perturbation risk:")
    add_formula(doc, "Rₜ,ₖ = (1 / |Iₜ|) Σᵢ∈Iₜ R(i,k)")
    doc.add_paragraph(
        "This metric identifies cases or perturbation types where semantic drift and correctness change occur together. "
        "It can be used to rank which perturbations create the highest robustness risk."
    )

    doc.add_heading("10. What Can Be Claimed as Innovation", level=1)
    doc.add_paragraph(
        "The safest wording is methodological extension, not a completely new theory."
    )
    doc.add_paragraph(
        "Suggested wording:"
    )
    doc.add_paragraph(
        "Building on sensitivity and consistency metrics for classification prompts, this study introduces a noise-corrected semantic sensitivity measure for generative LLM outputs. "
        "The measure subtracts within-prompt sampling variability from cross-prompt semantic drift, allowing perturbation-induced instability to be distinguished from stochastic generation noise. "
        "The study further links this semantic sensitivity to correctness drift, thereby extending prompt sensitivity analysis from output instability to functional robustness."
    )
    doc.add_paragraph("中文表述：")
    doc.add_paragraph(
        "本研究借鉴 classification prompt sensitivity/consistency 的思想，提出适用于生成式任务的 noise-corrected semantic sensitivity。"
        "该指标从 cross-prompt semantic drift 中扣除 within-prompt sampling variability，从而区分 perturbation-induced instability 和 stochastic generation noise。"
        "本研究进一步将该 semantic sensitivity 与 correctness drift 关联起来，把 prompt sensitivity analysis 从输出不稳定性扩展到功能正确性鲁棒性。"
    )

    doc.add_heading("11. Compact Formula Set for the Proposal", level=1)
    add_formula(doc, "D_raw(i,k) = (1 / R²) Σᵣ Σᵩ d(z⁰ᵢ,ᵣ, zᵏᵢ,ᵩ)")
    add_formula(doc, "N(i,k) = 1/2 {[2 / R(R - 1)]Σᵣ<ᵩ d(z⁰ᵢ,ᵣ, z⁰ᵢ,ᵩ) + [2 / R(R - 1)]Σᵣ<ᵩ d(zᵏᵢ,ᵣ, zᵏᵢ,ᵩ)}")
    add_formula(doc, "S_sem(i,k) = max(0, D_raw(i,k) - N(i,k))")
    add_formula(doc, "ΔC(i,k) = |C₀(i) - Cₖ(i)|")
    add_formula(doc, "ρ = Spearman(S_sem, ΔC)")
    add_formula(doc, "R(i,k) = S_sem(i,k) · ΔC(i,k)")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
