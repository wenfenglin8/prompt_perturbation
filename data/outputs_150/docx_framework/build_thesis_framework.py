from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/wenfenglin/Desktop/Pioneer/outputs/论文整体架构写作框架.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "D9E2EC"


def set_font(run, size=None, bold=None, italic=None, color=None, latin="Calibri", east_asia="宋体"):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_para_spacing(p, before=0, after=6, line=1.1):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_title_block(doc):
    p = doc.add_paragraph()
    set_para_spacing(p, after=4)
    r = p.add_run("论文整体架构与写作框架")
    set_font(r, size=23, bold=True, color=RGBColor(0, 0, 0), east_asia="黑体")

    p = doc.add_paragraph()
    set_para_spacing(p, after=12)
    r = p.add_run("主题：Prompt Robustness in Large Language Models")
    set_font(r, size=13, color=MUTED)

    rows = [
        ("用途", "作为论文正文写作模板；后续可直接在每个板块的“待填内容”下扩写。"),
        ("核心概念", "prompt robustness；sample noise；prompt perturbation；semantic drift；task type。"),
        ("研究问题", "RQ1：量化 sample noise 对输出稳定性的影响。RQ2：分析 prompt perturbations 对不同任务输出的影响。"),
        ("建议写法", "每节先写主张，再放方法/证据/图表，最后解释该节如何服务研究问题。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1600, 7760])
    for i, (label, text) in enumerate(rows):
        for j, value in enumerate([label, text]):
            cell = table.cell(i, j)
            shade_cell(cell, LIGHT_GRAY if j == 0 else "FFFFFF")
            p = cell.paragraphs[0]
            set_para_spacing(p, after=0)
            r = p.add_run(value)
            set_font(r, size=10.5, bold=(j == 0), color=RGBColor(0, 0, 0))


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para_spacing(p, after=4, line=1.167)
        r = p.add_run(item)
        set_font(r, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para_spacing(p, after=4, line=1.167)
        r = p.add_run(item)
        set_font(r, size=11)


def add_fill_box(doc, prompts, lines=2):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=4)
    r = p.add_run("待填内容：")
    set_font(r, size=10.5, bold=True, color=DARK_BLUE)
    for prompt in prompts:
        p = cell.add_paragraph()
        set_para_spacing(p, after=3, line=1.1)
        r = p.add_run(f"- {prompt}")
        set_font(r, size=10.5, color=RGBColor(0x22, 0x22, 0x22))
    for _ in range(lines):
        p = cell.add_paragraph()
        set_para_spacing(p, after=3)
        r = p.add_run("【在此扩写正文】")
        set_font(r, size=10.5, italic=True, color=MUTED)


def add_section(doc, title, objective, writing_points, fill_prompts, level=1):
    doc.add_heading(title, level=level)
    p = doc.add_paragraph()
    set_para_spacing(p, after=6)
    r = p.add_run("本节写作目标：")
    set_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(objective)
    set_font(r)
    add_bullets(doc, writing_points)
    add_fill_box(doc, fill_prompts)


def add_appendix_table(doc):
    doc.add_heading("附录与材料清单", level=1)
    p = doc.add_paragraph()
    set_para_spacing(p, after=6)
    r = p.add_run("建议在论文最后或研究文件夹中维护以下材料，写作时可直接回填到正文。")
    set_font(r)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_width(table, [1600, 2800, 2480, 2480])
    headers = ["材料", "用途", "当前状态", "后续补充"]
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        set_para_spacing(p, after=0)
        r = p.add_run(h)
        set_font(r, size=10.5, bold=True)
    rows = [
        ("原始 prompts", "说明样本来源与任务类型", "待整理", "补充筛选标准"),
        ("扰动 prompts", "支撑 RQ2 的 perturbation 分析", "待整理", "按扰动类型分类"),
        ("模型输出", "用于相似度与方差分析", "待整理", "保留 unaltered / perturbed 成对结果"),
        ("指标结果", "生成表格、tornado charts、heat charts", "待整理", "补充统计解释"),
        ("参考文献", "支撑任务类型、扰动类型、算法和方法选择", "待整理", "按 Introduction / Methodology / Analysis 分组"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            set_para_spacing(p, after=0)
            r = p.add_run(value)
            set_font(r, size=10.5)


def main():
    doc = Document()
    style_doc(doc)

    add_title_block(doc)

    add_section(
        doc,
        "摘要与关键词（写完正文后回填）",
        "用一段话概括研究背景、研究问题、方法、主要发现与意义；最后列出 4-6 个关键词。",
        [
            "摘要不要先写成最终版；建议正文完成后再压缩成 150-250 词左右。",
            "关键词可包括：Large Language Models, Prompt Robustness, Prompt Perturbation, Sample Noise, Semantic Similarity。",
        ],
        [
            "一句话说明研究关注：prompt 的微小变化是否会造成 LLM 输出不稳定。",
            "简述 RQ1 / RQ2、数据与实验方法、最重要的结果。",
            "最后指出该研究对日常学术使用 LLM 和 prompt 设计的意义。",
        ],
    )

    add_section(
        doc,
        "1. Introduction",
        "从个人观察和学术使用场景切入，解释为什么 prompt robustness 值得被单独研究。",
        [
            "以 personal engagement 或 brief data 开头：同一任务换一种表达，LLM 输出可能明显变化。",
            "说明 prompt 是模型输入，response 是模型输出，并自然引入输入扰动与输出变化的关系。",
            "定义 prompt robustness，并区分 robustness 与 accuracy：模型答得准不等于对提示词变化稳定。",
            "说明 sample noise 会让 single-generation result 具有偶然性，因此需要 baseline。",
            "说明 prompt perturbations 可能带来 semantic drift，不同 task types 与 perturbation types 可能影响不同。",
        ],
        [
            "写出你的个人观察或简短案例。",
            "补充 LLM 在日常 academic use 中的使用场景。",
            "插入 prompt robustness 的定义与引用。",
            "用一段话把 RQ1 和 RQ2 引出来。",
        ],
    )

    add_section(
        doc,
        "2. Research Questions",
        "明确列出研究问题，并解释每个问题在论文中的功能。",
        [
            "RQ1 聚焦 sample noise：即在没有刻意扰动 prompt 的情况下，模型输出本身会有多大波动。",
            "RQ2 聚焦 prompt perturbations：即不同扰动类型和任务类型如何影响输出语义相似度或文本重叠。",
            "每个 RQ 后可补一句预期贡献，说明它如何帮助更可靠地使用 LLM。",
        ],
        [
            "RQ1: To what extent does sample noise affect LLM output variance?",
            "RQ2: How do prompt perturbation types and task types affect the semantic drift of LLM outputs?",
            "解释 RQ1 是基线，RQ2 是在基线之上评估扰动效应。",
        ],
    )

    add_section(
        doc,
        "3. Literature Review / Conceptual Background",
        "集中安放后文需要依赖的概念和文献，避免 Introduction 过长。",
        [
            "整理 prompt engineering、prompt sensitivity、robustness、LLM evaluation 相关文献。",
            "为 task types、perturbation types、random sampling methodology、similarity metrics 寻找支撑。",
            "文献综述的结尾应指出 gap：已有研究可能关注 performance，但对 prompt robustness / sample-noise baseline 的处理不足。",
        ],
        [
            "列出 3-5 篇最关键文献及其功能。",
            "写出本文采用哪些概念定义，哪些定义需要与 previous research 对齐。",
            "用一段 gap statement 过渡到 Methodology。",
        ],
    )

    add_section(
        doc,
        "4. Methodology",
        "解释研究设计、模型选择、数据集选择、prompt 扰动生成、输出生成和相似度测量方法。",
        [
            "说明为什么选择两个 LLM types：一个用于生成 perturbations，另一个用于输出 generation。",
            "说明 datasets 的选择理由，以及它们如何覆盖不同 task types。",
            "写清楚 RQ1 的 sample-noise measurement：随机抽样、重复生成、方差/相似度指标。",
            "写清楚 RQ2 的 perturbation workflow：原始 prompt、扰动 prompt、输出生成、paired comparison。",
            "说明 similarity measures：Sentence-BERT embeddings、text-overlap metrics，或其他你最终采用的指标。",
        ],
        [
            "模型选择理由与引用。",
            "数据集与 task types 的定义。",
            "perturbation types 的定义、生成方法和例子。",
            "RQ1 / RQ2 的算法流程或伪代码。",
            "指标公式、计算工具和统计方法。",
        ],
    )

    doc.add_heading("4.1 RQ1: Sample Noise Baseline", level=2)
    add_numbered(
        doc,
        [
            "从不同 task types 中随机抽取 prompts。",
            "在不改变 prompt 的前提下重复生成 outputs。",
            "计算 outputs 之间的语义相似度、文本重叠或其他稳定性指标。",
            "将该结果作为后续判断 perturbation effect 的 baseline。",
        ],
    )
    add_fill_box(doc, ["写明 sample noise 的操作性定义。", "补充参考文献中类似 methodology 的做法。", "说明为什么 single-generation results may be misleading。"])

    doc.add_heading("4.2 RQ2: Prompt Perturbation Effect", level=2)
    add_numbered(
        doc,
        [
            "选择 task types 与 perturbation types，并说明选择理由。",
            "使用第一个 LLM 生成扰动后的 prompts。",
            "使用第二个 LLM 分别生成 unaltered prompt output 与 perturbed prompt outputs。",
            "对成对 outputs 计算 similarity / overlap，并分析 task type 与 perturbation type 的交互关系。",
        ],
    )
    add_fill_box(doc, ["列出 perturbation types 的最终分类。", "补充每种 perturbation 的例子。", "说明如何控制随机性、温度或重复次数。"])

    add_section(
        doc,
        "5. Experiments",
        "把实验设置写成可复现流程，读者应能根据本节复刻你的数据生成与分析步骤。",
        [
            "Set up：模型版本、参数、prompt 数量、抽样规则、运行次数。",
            "Environment：运行环境、代码结构、主要库、硬件/平台信息。",
            "Procedure：按步骤说明从 prompt sampling 到 result visualization 的全流程。",
            "建议在本节加入流程图或步骤列表，减少正文解释压力。",
        ],
        [
            "填写模型参数与实验环境。",
            "填写每个 task type 的样本量。",
            "填写 output collection 与 data cleaning 规则。",
            "填写实验流程编号列表。",
        ],
    )

    add_section(
        doc,
        "6. Results and Analysis",
        "围绕 RQ1 与 RQ2 分别呈现结果，再解释这些结果说明了什么。",
        [
            "RQ1：报告 sample noise 导致的 output variance，并与 previous research 比较。",
            "RQ2：分析 task types 与 perturbation types 的关系，指出最 potent perturbation type。",
            "解释为什么某些扰动更容易改变输出，例如语义范围变化、任务约束变化、关键词替换等。",
            "使用 tornado charts 可展示不同扰动类型的影响强度；使用 heat charts 展示 task type 与 perturbation type 的交互。",
        ],
        [
            "RQ1 的主要数值结果和表格编号。",
            "RQ2 的主要数值结果和图表编号。",
            "写出 2-3 个最重要发现，每个发现后接解释。",
            "标注需要插入 tornado chart / heat chart 的位置。",
        ],
    )

    add_section(
        doc,
        "7. Evaluation / Discussion",
        "把结果提升为讨论：解释结果的可靠性、意义、与研究问题的对应关系。",
        [
            "讨论指标是否足以捕捉 semantic drift，哪些结果可能只是 sample noise。",
            "说明研究对 prompt design、LLM evaluation、academic use of LLM 的启示。",
            "把 RQ1 baseline 和 RQ2 perturbation results 连接起来，避免两个实验像分开的项目。",
        ],
        [
            "评价实验设计的有效性。",
            "说明结果如何回答每个 RQ。",
            "写出 practical implications 和 theoretical implications。",
        ],
    )

    add_section(
        doc,
        "8. Limitations and Future Improvements",
        "诚实说明研究边界，并提出后续可改进方向。",
        [
            "可能限制：模型数量有限、数据集范围有限、扰动生成依赖另一个 LLM、指标无法完全代表人类判断。",
            "改进方向：扩大模型与任务类型、加入人工评估、加入更多语言/学科场景、测试更多 prompt perturbation categories。",
            "避免只列限制；每个限制后最好配一个可执行的 future improvement。",
        ],
        [
            "列出 3-5 个 limitations。",
            "为每个 limitation 写出对应 improvement。",
            "最后用一句话回到 prompt robustness 的研究价值。",
        ],
    )

    add_section(
        doc,
        "9. Conclusion",
        "简洁回收全文：研究做了什么、发现了什么、为什么重要。",
        [
            "不要引入新数据或新文献。",
            "按 Introduction 中提出的问题顺序回应 RQ1 和 RQ2。",
            "结尾可强调：理解 prompt robustness 有助于更可靠地解释和使用 LLM outputs。",
        ],
        [
            "一句话总结研究目的。",
            "两三句话总结主要发现。",
            "一句话总结贡献与未来方向。",
        ],
    )

    add_appendix_table(doc)

    doc.add_heading("References（后续按学校格式整理）", level=1)
    p = doc.add_paragraph()
    set_para_spacing(p, after=6)
    r = p.add_run("【在此添加参考文献。建议先按主题分组：prompt robustness / sample noise / prompt perturbation / similarity metrics / methodology。】")
    set_font(r, italic=True, color=MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
